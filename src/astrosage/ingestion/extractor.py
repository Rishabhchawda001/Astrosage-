"""
Adaptive document extraction pipeline.

Routing strategy:
  1. Native PDF → PyMuPDF (fast, lossless)
  2. Scanned/mixed PDF → PaddleOCR
  3. Complex layouts → Docling
  4. Non-PDF → format-specific handlers
"""
from __future__ import annotations

import logging
from pathlib import Path

from ..models import (
    Document,
    DocumentMetadata,
    DocumentType,
    OCRStatus,
    PageContent,
    Script,
    generate_id,
    now_utc,
)

logger = logging.getLogger(__name__)


def detect_document_type(filepath: Path) -> DocumentType:
    ext = filepath.suffix.lower()
    mapping = {
        ".pdf": DocumentType.PDF,
        ".docx": DocumentType.DOCX,
        ".doc": DocumentType.DOCX,
        ".epub": DocumentType.EPUB,
        ".txt": DocumentType.TXT,
        ".md": DocumentType.MARKDOWN,
        ".jpg": DocumentType.IMAGE,
        ".jpeg": DocumentType.IMAGE,
        ".png": DocumentType.IMAGE,
        ".gif": DocumentType.IMAGE,
        ".tiff": DocumentType.IMAGE,
        ".zip": DocumentType.ZIP,
    }
    return mapping.get(ext, DocumentType.UNKNOWN)


def detect_script(text: str) -> Script:
    devanagari_count = sum(1 for c in text if "\u0900" <= c <= "\u097F")
    latin_count = sum(1 for c in text if c.isascii() and c.isalpha())
    total = devanagari_count + latin_count

    if total == 0:
        return Script.UNKNOWN

    devanagari_ratio = devanagari_count / total
    if devanagari_ratio > 0.7:
        return Script.DEVANAGARI
    elif devanagari_ratio < 0.3:
        return Script.LATIN
    else:
        return Script.MIXED


def _is_native_pdf(filepath: Path) -> bool:
    """Check if PDF contains extractable text (not scanned)."""
    try:
        import pymupdf

        doc = pymupdf.open(str(filepath))
        total_chars = 0
        sample_pages = min(5, len(doc))
        for i in range(sample_pages):
            page = doc[i]
            total_chars += len(page.get_text())
        doc.close()
        return total_chars > 100
    except Exception:
        return False


def extract_native_pdf(filepath: Path) -> Document:
    """Extract text from native (text-based) PDF using PyMuPDF."""
    import pymupdf

    logger.info(f"Extracting native PDF: {filepath.name}")

    doc_pymupdf = pymupdf.open(str(filepath))
    pages = []

    for i in range(len(doc_pymupdf)):
        page = doc_pymupdf[i]
        text = page.get_text()
        has_images = len(page.get_images()) > 0
        has_tables = False  # pymupdf doesn't directly detect tables

        pages.append(
            PageContent(
                page_number=i + 1,
                text=text.strip(),
                has_images=has_images,
                has_tables=has_tables,
                ocr_confidence=1.0,
                ocr_engine="pymupdf",
                language_detected=detect_script(text).value,
            )
        )

    page_count = len(doc_pymupdf)
    doc_pymupdf.close()

    metadata = DocumentMetadata(
        document_id=generate_id("doc"),
        source_filename=filepath.name,
        source_path=str(filepath),
        document_type=DocumentType.PDF,
        ocr_status=OCRStatus.NOT_NEEDED,
        page_count=page_count,
        file_size_bytes=filepath.stat().st_size,
    )

    return Document(metadata=metadata, pages=pages)


def extract_scanned_pdf(filepath: Path) -> Document:
    """Extract text from scanned PDF using PaddleOCR."""
    logger.info(f"OCR processing: {filepath.name}")

    try:
        from paddleocr import PaddleOCR
    except ImportError:
        logger.error("PaddleOCR not installed. Falling back to Tesseract.")
        return extract_with_tesseract(filepath)

    # Detect language from filename/path
    path_str = str(filepath).lower()
    if any(kw in path_str for kw in ["english", "eng"]):
        lang = "en"
    elif any(kw in path_str for kw in ["hindi", "hind"]):
        lang = "hi"
    elif any(kw in path_str for kw in ["sanskrit", "san"]):
        lang = "sa"
    else:
        lang = "en"  # Default to English

    ocr = PaddleOCR(use_angle_cls=True, lang=lang, show_log=False)

    import pymupdf

    doc_pymupdf = pymupdf.open(str(filepath))
    pages = []

    for i in range(len(doc_pymupdf)):
        page = doc_pymupdf[i]

        # Convert page to image for OCR
        pix = page.get_pixmap(dpi=200)
        img_bytes = pix.tobytes("png")

        # Save temporarily
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(img_bytes)
            tmp_path = tmp.name

        try:
            result = ocr.ocr(tmp_path, cls=True)
            text_parts = []
            avg_confidence = 0.0

            if result and result[0]:
                for line in result[0]:
                    if line[1]:
                        text_parts.append(line[1][0])
                        avg_confidence += line[1][1]

                if len(result[0]) > 0:
                    avg_confidence /= len(result[0])

            text = "\n".join(text_parts)
        finally:
            Path(tmp_path).unlink(missing_ok=True)

        pages.append(
            PageContent(
                page_number=i + 1,
                text=text.strip(),
                has_images=True,
                ocr_confidence=avg_confidence,
                ocr_engine="paddleocr",
                language_detected=lang,
            )
        )

    page_count = len(doc_pymupdf)
    doc_pymupdf.close()

    min_confidence = min(p.ocr_confidence for p in pages) if pages else 0
    ocr_status = (
        OCRStatus.COMPLETED
        if min_confidence > 0.7
        else OCRStatus.REVIEW_NEEDED
    )

    metadata = DocumentMetadata(
        document_id=generate_id("doc"),
        source_filename=filepath.name,
        source_path=str(filepath),
        document_type=DocumentType.PDF,
        ocr_status=ocr_status,
        ocr_confidence=min_confidence,
        ocr_engine="paddleocr",
        page_count=page_count,
        file_size_bytes=filepath.stat().st_size,
    )

    return Document(metadata=metadata, pages=pages)


def extract_with_tesseract(filepath: Path) -> Document:
    """Fallback OCR using Tesseract via pytesseract."""
    import tempfile

    try:
        import pytesseract
        import pymupdf
    except ImportError as e:
        logger.error(f"Missing dependency: {e}")
        raise

    logger.info(f"Tesseract OCR: {filepath.name}")

    doc_pymupdf = pymupdf.open(str(filepath))
    pages = []

    for i in range(len(doc_pymupdf)):
        page = doc_pymupdf[i]
        pix = page.get_pixmap(dpi=200)

        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            pix.save(tmp.name)
            text = pytesseract.image_to_string(tmp.name)
            confidence_data = pytesseract.image_to_data(
                tmp.name, output_type=pytesseract.Output.DICT
            )
            confs = [
                int(c)
                for c in confidence_data["conf"]
                if isinstance(c, (int, float)) and int(c) > 0
            ]
            avg_conf = sum(confs) / len(confs) if confs else 0

        Path(tmp.name).unlink(missing_ok=True)

        pages.append(
            PageContent(
                page_number=i + 1,
                text=text.strip(),
                ocr_confidence=avg_conf / 100.0,
                ocr_engine="tesseract",
            )
        )

    page_count = len(doc_pymupdf)
    doc_pymupdf.close()

    metadata = DocumentMetadata(
        document_id=generate_id("doc"),
        source_filename=filepath.name,
        source_path=str(filepath),
        document_type=DocumentType.PDF,
        ocr_status=OCRStatus.COMPLETED,
        ocr_engine="tesseract",
        page_count=page_count,
        file_size_bytes=filepath.stat().st_size,
    )

    return Document(metadata=metadata, pages=pages)


def extract_docx(filepath: Path) -> Document:
    """Extract text from DOCX files."""
    from docx import Document as DocxDocument

    logger.info(f"Extracting DOCX: {filepath.name}")

    doc = DocxDocument(str(filepath))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    full_text = "\n\n".join(paragraphs)
    script = detect_script(full_text)

    metadata = DocumentMetadata(
        document_id=generate_id("doc"),
        source_filename=filepath.name,
        source_path=str(filepath),
        document_type=DocumentType.DOCX,
        ocr_status=OCRStatus.NOT_NEEDED,
        file_size_bytes=filepath.stat().st_size,
        script=script,
    )

    page = PageContent(
        page_number=1,
        text=full_text,
        language_detected=script.value,
    )

    return Document(metadata=metadata, pages=[page])


def extract_text_file(filepath: Path) -> Document:
    """Extract from plain text or markdown."""
    text = filepath.read_text(encoding="utf-8", errors="replace")
    script = detect_script(text)

    ext = filepath.suffix.lower()
    doc_type = DocumentType.MARKDOWN if ext == ".md" else DocumentType.TXT

    metadata = DocumentMetadata(
        document_id=generate_id("doc"),
        source_filename=filepath.name,
        source_path=str(filepath),
        document_type=doc_type,
        ocr_status=OCRStatus.NOT_NEEDED,
        file_size_bytes=filepath.stat().st_size,
        script=script,
    )

    page = PageContent(
        page_number=1,
        text=text,
        language_detected=script.value,
    )

    return Document(metadata=metadata, pages=[page])


def extract_document(filepath: Path) -> Document:
    """
    Adaptive extraction: routes to the appropriate extractor based on
    document type and content analysis.
    """
    doc_type = detect_document_type(filepath)

    if doc_type == DocumentType.PDF:
        if _is_native_pdf(filepath):
            return extract_native_pdf(filepath)
        else:
            return extract_scanned_pdf(filepath)
    elif doc_type == DocumentType.DOCX:
        return extract_docx(filepath)
    elif doc_type in (DocumentType.TXT, DocumentType.MARKDOWN):
        return extract_text_file(filepath)
    elif doc_type == DocumentType.IMAGE:
        return extract_scanned_pdf(filepath)
    else:
        logger.warning(f"Unsupported document type: {doc_type} for {filepath.name}")
        metadata = DocumentMetadata(
            document_id=generate_id("doc"),
            source_filename=filepath.name,
            source_path=str(filepath),
            document_type=doc_type,
            ocr_status=OCRStatus.FAILED,
            file_size_bytes=filepath.stat().st_size,
        )
        return Document(metadata=metadata, pages=[])
